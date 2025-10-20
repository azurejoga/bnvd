"""
Sistema de Exportação de Dados do BNVD
Permite exportar vulnerabilidades e notícias em múltiplos formatos
"""

import os
import csv
import json
import logging
from io import BytesIO, StringIO
from datetime import datetime, timezone, timedelta
from typing import Dict, List, Any, Optional
from flask import Response, make_response
import xml.etree.ElementTree as ET
from xml.dom import minidom

# Tentativa de importar bibliotecas opcionais
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False
    logging.warning("reportlab não está instalado - exportação PDF desabilitada")

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logging.warning("python-docx não está instalado - exportação DOCX desabilitada")

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def get_brazil_time():
    """Retorna a hora atual no fuso horário do Brasil (UTC-3)"""
    utc_now = datetime.now(timezone.utc)
    brazil_tz = timezone(timedelta(hours=-3))
    return utc_now.astimezone(brazil_tz)

class ExportManager:
    """Gerenciador de exportações do BNVD"""

    def __init__(self):
        self.supported_formats = ['csv', 'json', 'pdf', 'txt', 'docx', 'xml', 'odf']

    def export_vulnerability(self, vuln_data: Dict, format_type: str) -> Response:
        """
        Exporta uma vulnerabilidade no formato especificado

        Args:
            vuln_data: Dados da vulnerabilidade
            format_type: Formato de exportação (csv, json, pdf, txt, docx, xml, odf)

        Returns:
            Flask Response com o arquivo para download
        """
        if format_type not in self.supported_formats:
            raise ValueError(f"Formato não suportado: {format_type}")

        cve_id = vuln_data.get('cve', {}).get('id', 'unknown')
        filename = f"{cve_id}_{get_brazil_time().strftime('%Y%m%d_%H%M%S')}"

        if format_type == 'csv':
            return self._export_vuln_csv(vuln_data, filename)
        elif format_type == 'json':
            return self._export_vuln_json(vuln_data, filename)
        elif format_type == 'pdf':
            return self._export_vuln_pdf(vuln_data, filename)
        elif format_type == 'txt':
            return self._export_vuln_txt(vuln_data, filename)
        elif format_type == 'docx':
            return self._export_vuln_docx(vuln_data, filename)
        elif format_type == 'xml':
            return self._export_vuln_xml(vuln_data, filename)
        elif format_type == 'odf':
            return self._export_vuln_odf(vuln_data, filename)
        else:
            # Fallback para TXT se nenhum formato corresponder
            return self._export_vuln_txt(vuln_data, filename)

    def export_news(self, news_data: Dict, format_type: str) -> Response:
        """
        Exporta uma notícia no formato especificado

        Args:
            news_data: Dados da notícia
            format_type: Formato de exportação

        Returns:
            Flask Response com o arquivo para download
        """
        if format_type not in self.supported_formats:
            raise ValueError(f"Formato não suportado: {format_type}")

        slug = news_data.get('slug', 'noticia')
        filename = f"noticia_{slug}_{get_brazil_time().strftime('%Y%m%d_%H%M%S')}"

        if format_type == 'csv':
            return self._export_news_csv(news_data, filename)
        elif format_type == 'json':
            return self._export_news_json(news_data, filename)
        elif format_type == 'pdf':
            return self._export_news_pdf(news_data, filename)
        elif format_type == 'txt':
            return self._export_news_txt(news_data, filename)
        elif format_type == 'docx':
            return self._export_news_docx(news_data, filename)
        elif format_type == 'xml':
            return self._export_news_xml(news_data, filename)
        elif format_type == 'odf':
            return self._export_news_odf(news_data, filename)
        else:
            # Fallback para TXT se nenhum formato corresponder
            return self._export_news_txt(news_data, filename)

    # ==================== Exportação de Vulnerabilidades ====================

    def _export_vuln_csv(self, vuln_data: Dict, filename: str) -> Response:
        """Exporta vulnerabilidade para CSV com todos os campos no formato padronizado"""
        si = StringIO()
        writer = csv.writer(si)

        cve = vuln_data.get('cve', {})
        cve_id = cve.get('id', 'CVE Desconhecido')

        # Cabeçalho
        writer.writerow(['BNVD - Banco Nacional de Vulnerabilidades Cibernéticas'])
        writer.writerow([''])
        writer.writerow(['Campo', 'Valor'])
        writer.writerow([''])

        # Dados básicos
        writer.writerow(['CVE ID', cve_id])
        writer.writerow(['Status', cve.get('vulnStatus', 'N/A')])
        writer.writerow(['Publicado', cve.get('published', 'N/A')])
        writer.writerow(['Última Modificação', cve.get('lastModified', 'N/A')])
        writer.writerow(['Fonte', cve.get('sourceIdentifier', 'N/A')])
        writer.writerow([''])

        # Descrições
        writer.writerow(['--- DESCRIÇÃO ORIGINAL (INGLÊS) ---'])
        descriptions = cve.get('descriptions', [])
        for desc in descriptions:
            if desc.get('lang') == 'en':
                writer.writerow(['', desc.get('value', 'N/A')])
                break
        writer.writerow([''])

        writer.writerow(['--- DESCRIÇÃO TRADUZIDA (PORTUGUÊS) ---'])
        for desc in descriptions:
            if desc.get('lang') == 'en' and desc.get('value_pt'):
                writer.writerow(['', desc.get('value_pt', 'N/A')])
                break
        writer.writerow([''])

        # CVSS Metrics completo
        writer.writerow(['--- MÉTRICAS CVSS DETALHADAS ---'])
        metrics = cve.get('metrics', {})
        for version in ['cvssMetricV31', 'cvssMetricV30', 'cvssMetricV2']:
            if version in metrics and metrics[version]:
                metric = metrics[version][0]
                cvss_data = metric.get('cvssData', {})

                writer.writerow([''])
                writer.writerow([f'Versão', version])
                writer.writerow(['Vector String', cvss_data.get('vectorString', 'N/A')])
                writer.writerow(['Score Base', cvss_data.get('baseScore', 'N/A')])
                writer.writerow(['Severidade', cvss_data.get('baseSeverity', 'N/A')])
                writer.writerow(['Vetor de Ataque', cvss_data.get('attackVector', 'N/A')])
                writer.writerow(['Complexidade', cvss_data.get('attackComplexity', 'N/A')])
                writer.writerow(['Privilégios Necessários', cvss_data.get('privilegesRequired', 'N/A')])
                writer.writerow(['Interação do Usuário', cvss_data.get('userInteraction', 'N/A')])
                writer.writerow(['Escopo', cvss_data.get('scope', 'N/A')])
                writer.writerow(['Impacto Confidencialidade', cvss_data.get('confidentialityImpact', 'N/A')])
                writer.writerow(['Impacto Integridade', cvss_data.get('integrityImpact', 'N/A')])
                writer.writerow(['Impacto Disponibilidade', cvss_data.get('availabilityImpact', 'N/A')])
                break

        # CWE
        weaknesses = cve.get('weaknesses', [])
        if weaknesses:
            writer.writerow([''])
            writer.writerow(['--- TIPOS DE FRAQUEZA (CWE) ---'])
            for weakness in weaknesses:
                for desc in weakness.get('description', []):
                    writer.writerow(['CWE', desc.get('value', 'N/A')])

        # Referências
        references = cve.get('references', [])
        if references:
            writer.writerow([''])
            writer.writerow(['--- REFERÊNCIAS ---'])
            for ref in references:
                writer.writerow([''])
                writer.writerow(['URL', ref.get('url', 'N/A')])
                writer.writerow(['Fonte', ref.get('source', 'N/A')])
                tags = ref.get('tags', [])
                if tags:
                    writer.writerow(['Tags', ', '.join(tags)])

        writer.writerow([''])
        writer.writerow(['Exportado em', get_brazil_time().strftime('%d/%m/%Y às %H:%M:%S') + ' (Horário de Brasília)'])

        output = make_response(si.getvalue())
        output.headers["Content-Disposition"] = f"attachment; filename={filename}.csv"
        output.headers["Content-type"] = "text/csv; charset=utf-8"
        return output

    def _export_vuln_json(self, vuln_data: Dict, filename: str) -> Response:
        """Exporta vulnerabilidade para JSON com formato padronizado"""
        cve = vuln_data.get('cve', {})
        cve_id = cve.get('id', 'CVE Desconhecido')

        # Estrutura padronizada seguindo modelo TXT
        export_data = {
            'bnvd': 'Banco Nacional de Vulnerabilidades Cibernéticas',
            'tipo': 'Vulnerabilidade',
            'cve_id': cve_id,
            'status': cve.get('vulnStatus', 'N/A'),
            'publicado': cve.get('published', 'N/A'),
            'ultima_modificacao': cve.get('lastModified', 'N/A'),
            'fonte': cve.get('sourceIdentifier', 'N/A'),
            'descricao_original': '',
            'descricao_traduzida': '',
            'metricas_cvss': {},
            'tipos_fraqueza': [],
            'referencias': [],
            'exportado_em': get_brazil_time().strftime('%d/%m/%Y às %H:%M:%S') + ' (Horário de Brasília)'
        }

        # Descrições
        descriptions = cve.get('descriptions', [])
        for desc in descriptions:
            if desc.get('lang') == 'en':
                export_data['descricao_original'] = desc.get('value', 'N/A')
                export_data['descricao_traduzida'] = desc.get('value_pt', desc.get('value', 'N/A'))
                break

        # CVSS Metrics
        metrics = cve.get('metrics', {})
        for version in ['cvssMetricV31', 'cvssMetricV30', 'cvssMetricV2']:
            if version in metrics and metrics[version]:
                metric = metrics[version][0]
                cvss_data = metric.get('cvssData', {})
                export_data['metricas_cvss'] = {
                    'versao': version,
                    'vector_string': cvss_data.get('vectorString', 'N/A'),
                    'score_base': cvss_data.get('baseScore', 'N/A'),
                    'severidade': cvss_data.get('baseSeverity', 'N/A'),
                    'vetor_ataque': cvss_data.get('attackVector', 'N/A'),
                    'complexidade': cvss_data.get('attackComplexity', 'N/A'),
                    'privilegios_necessarios': cvss_data.get('privilegesRequired', 'N/A'),
                    'interacao_usuario': cvss_data.get('userInteraction', 'N/A'),
                    'escopo': cvss_data.get('scope', 'N/A'),
                    'impacto_confidencialidade': cvss_data.get('confidentialityImpact', 'N/A'),
                    'impacto_integridade': cvss_data.get('integrityImpact', 'N/A'),
                    'impacto_disponibilidade': cvss_data.get('availabilityImpact', 'N/A')
                }
                break

        # CWE
        weaknesses = cve.get('weaknesses', [])
        for weakness in weaknesses:
            for desc in weakness.get('description', []):
                export_data['tipos_fraqueza'].append(desc.get('value', 'N/A'))

        # Referências
        references = cve.get('references', [])
        for ref in references:
            export_data['referencias'].append({
                'url': ref.get('url', 'N/A'),
                'fonte': ref.get('source', 'N/A'),
                'tags': ref.get('tags', [])
            })

        output = make_response(json.dumps(export_data, indent=2, ensure_ascii=False))
        output.headers["Content-Disposition"] = f"attachment; filename={filename}.json"
        output.headers["Content-type"] = "application/json; charset=utf-8"
        return output

    def _export_vuln_pdf(self, vuln_data: Dict, filename: str) -> Response:
        """Exporta vulnerabilidade para PDF no formato padronizado"""
        if not HAS_REPORTLAB:
            logging.warning("ReportLab não disponível, usando formato TXT como fallback para PDF")
            return self._export_vuln_txt(vuln_data, filename.replace('.pdf', '.txt'))

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        story = []
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=20, textColor=colors.HexColor('#0d6efd'), spaceAfter=20, alignment=TA_CENTER)
        heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading2'], fontSize=14, textColor=colors.HexColor('#0d6efd'), spaceAfter=10)

        cve = vuln_data.get('cve', {})
        cve_id = cve.get('id', 'CVE Desconhecido')

        # Cabeçalho BNVD
        story.append(Paragraph('BNVD - Banco Nacional de Vulnerabilidades Cibernéticas', heading_style))
        story.append(Spacer(1, 12))
        story.append(Paragraph(cve_id, title_style))
        story.append(Spacer(1, 12))

        # Dados básicos
        story.append(Paragraph(f"<b>Status:</b> {cve.get('vulnStatus', 'N/A')}", styles['Normal']))
        story.append(Paragraph(f"<b>Publicado:</b> {cve.get('published', 'N/A')}", styles['Normal']))
        story.append(Paragraph(f"<b>Última Modificação:</b> {cve.get('lastModified', 'N/A')}", styles['Normal']))
        story.append(Paragraph(f"<b>Fonte:</b> {cve.get('sourceIdentifier', 'N/A')}", styles['Normal']))
        story.append(Spacer(1, 20))

        # Descrições
        story.append(Paragraph('DESCRIÇÃO ORIGINAL (INGLÊS)', heading_style))
        descriptions = cve.get('descriptions', [])
        for desc in descriptions:
            if desc.get('lang') == 'en':
                story.append(Paragraph(desc.get('value', 'N/A'), styles['Normal']))
                break
        story.append(Spacer(1, 12))

        story.append(Paragraph('DESCRIÇÃO TRADUZIDA (PORTUGUÊS)', heading_style))
        for desc in descriptions:
            if desc.get('lang') == 'en' and desc.get('value_pt'):
                story.append(Paragraph(desc.get('value_pt', 'N/A'), styles['Normal']))
                break
        story.append(Spacer(1, 20))

        # CVSS Metrics Detalhado
        story.append(Paragraph('MÉTRICAS CVSS DETALHADAS', heading_style))
        metrics = cve.get('metrics', {})
        for version in ['cvssMetricV31', 'cvssMetricV30', 'cvssMetricV2']:
            if version in metrics and metrics[version]:
                metric = metrics[version][0]
                cvss_data = metric.get('cvssData', {})
                story.append(Paragraph(f"<b>Versão:</b> {version}", styles['Normal']))
                story.append(Paragraph(f"<b>Vector String:</b> {cvss_data.get('vectorString', 'N/A')}", styles['Normal']))
                story.append(Paragraph(f"<b>Score Base:</b> {cvss_data.get('baseScore', 'N/A')}", styles['Normal']))
                story.append(Paragraph(f"<b>Severidade:</b> {cvss_data.get('baseSeverity', 'N/A')}", styles['Normal']))
                story.append(Paragraph(f"<b>Vetor de Ataque:</b> {cvss_data.get('attackVector', 'N/A')}", styles['Normal']))
                story.append(Paragraph(f"<b>Complexidade:</b> {cvss_data.get('attackComplexity', 'N/A')}", styles['Normal']))
                story.append(Paragraph(f"<b>Privilégios Necessários:</b> {cvss_data.get('privilegesRequired', 'N/A')}", styles['Normal']))
                story.append(Paragraph(f"<b>Interação do Usuário:</b> {cvss_data.get('userInteraction', 'N/A')}", styles['Normal']))
                story.append(Paragraph(f"<b>Escopo:</b> {cvss_data.get('scope', 'N/A')}", styles['Normal']))
                story.append(Paragraph(f"<b>Impacto Confidencialidade:</b> {cvss_data.get('confidentialityImpact', 'N/A')}", styles['Normal']))
                story.append(Paragraph(f"<b>Impacto Integridade:</b> {cvss_data.get('integrityImpact', 'N/A')}", styles['Normal']))
                story.append(Paragraph(f"<b>Impacto Disponibilidade:</b> {cvss_data.get('availabilityImpact', 'N/A')}", styles['Normal']))
                break
        story.append(Spacer(1, 20))

        # CWE
        weaknesses = cve.get('weaknesses', [])
        if weaknesses:
            story.append(Paragraph('TIPOS DE FRAQUEZA (CWE)', heading_style))
            for weakness in weaknesses:
                for desc in weakness.get('description', []):
                    story.append(Paragraph(f"• {desc.get('value', 'N/A')}", styles['Normal']))
        story.append(Spacer(1, 20))

        # Referências
        references = cve.get('references', [])
        if references:
            story.append(Paragraph('REFERÊNCIAS', heading_style))
            for ref in references:
                story.append(Paragraph(f"<b>URL:</b> {ref.get('url', 'N/A')}", styles['Normal']))
                story.append(Paragraph(f"<b>Fonte:</b> {ref.get('source', 'N/A')}", styles['Normal']))
                tags = ref.get('tags', [])
                if tags:
                    story.append(Paragraph(f"<b>Tags:</b> {', '.join(tags)}", styles['Normal']))
                story.append(Spacer(1, 6))

        story.append(Spacer(1, 20))
        story.append(Paragraph(f"Exportado em: {get_brazil_time().strftime('%d/%m/%Y às %H:%M:%S')} (Horário de Brasília)", styles['Normal']))

        doc.build(story)
        pdf_data = buffer.getvalue()
        buffer.close()

        output = make_response(pdf_data)
        output.headers["Content-Disposition"] = f"attachment; filename={filename}.pdf"
        output.headers["Content-type"] = "application/pdf"
        return output

    def _export_vuln_txt(self, vuln_data: Dict, filename: str) -> Response:
        """Exporta vulnerabilidade para TXT com todos os campos"""
        cve = vuln_data.get('cve', {})
        cve_id = cve.get('id', 'CVE Desconhecido')

        content = []
        content.append("=" * 80)
        content.append(f"BNVD - Banco Nacional de Vulnerabilidades Cibernéticas")
        content.append("=" * 80)
        content.append("")
        content.append(f"CVE ID: {cve_id}")
        content.append(f"Status: {cve.get('vulnStatus', 'N/A')}")
        content.append(f"Publicado: {cve.get('published', 'N/A')}")
        content.append(f"Última Modificação: {cve.get('lastModified', 'N/A')}")
        content.append(f"Fonte: {cve.get('sourceIdentifier', 'N/A')}")

        content.append("")
        content.append("-" * 80)
        content.append("DESCRIÇÃO ORIGINAL (INGLÊS)")
        content.append("-" * 80)
        descriptions = cve.get('descriptions', [])
        for desc in descriptions:
            if desc.get('lang') == 'en':
                content.append(desc.get('value', 'N/A'))
                break

        content.append("")
        content.append("-" * 80)
        content.append("DESCRIÇÃO TRADUZIDA (PORTUGUÊS)")
        content.append("-" * 80)
        for desc in descriptions:
            if desc.get('lang') == 'en' and desc.get('value_pt'):
                content.append(desc.get('value_pt', 'N/A'))
                break

        content.append("")
        content.append("-" * 80)
        content.append("MÉTRICAS CVSS DETALHADAS")
        content.append("-" * 80)

        metrics = cve.get('metrics', {})
        for version in ['cvssMetricV31', 'cvssMetricV30', 'cvssMetricV2']:
            if version in metrics and metrics[version]:
                metric = metrics[version][0]
                cvss_data = metric.get('cvssData', {})
                content.append(f"\nVersão: {version}")
                content.append(f"  Vector String: {cvss_data.get('vectorString', 'N/A')}")
                content.append(f"  Score Base: {cvss_data.get('baseScore', 'N/A')}")
                content.append(f"  Severidade: {cvss_data.get('baseSeverity', 'N/A')}")
                content.append(f"  Vetor de Ataque: {cvss_data.get('attackVector', 'N/A')}")
                content.append(f"  Complexidade: {cvss_data.get('attackComplexity', 'N/A')}")
                content.append(f"  Privilégios Necessários: {cvss_data.get('privilegesRequired', 'N/A')}")
                content.append(f"  Interação do Usuário: {cvss_data.get('userInteraction', 'N/A')}")
                content.append(f"  Escopo: {cvss_data.get('scope', 'N/A')}")
                content.append(f"  Impacto Confidencialidade: {cvss_data.get('confidentialityImpact', 'N/A')}")
                content.append(f"  Impacto Integridade: {cvss_data.get('integrityImpact', 'N/A')}")
                content.append(f"  Impacto Disponibilidade: {cvss_data.get('availabilityImpact', 'N/A')}")
                break

        # CWE
        weaknesses = cve.get('weaknesses', [])
        if weaknesses:
            content.append("")
            content.append("-" * 80)
            content.append("TIPOS DE FRAQUEZA (CWE)")
            content.append("-" * 80)
            for weakness in weaknesses:
                for desc in weakness.get('description', []):
                    content.append(f"• {desc.get('value', 'N/A')}")

        references = cve.get('references', [])
        if references:
            content.append("")
            content.append("-" * 80)
            content.append("REFERÊNCIAS")
            content.append("-" * 80)
            for ref in references:
                content.append(f"\nURL: {ref.get('url', '')}")
                content.append(f"Fonte: {ref.get('source', 'N/A')}")
                tags = ref.get('tags', [])
                if tags:
                    content.append(f"Tags: {', '.join(tags)}")

        content.append("")
        content.append("=" * 80)
        content.append(f"Exportado em: {get_brazil_time().strftime('%d/%m/%Y às %H:%M:%S')} (Horário de Brasília)")
        content.append("=" * 80)

        txt_data = "\n".join(content)
        output = make_response(txt_data)
        output.headers["Content-Disposition"] = f"attachment; filename={filename}.txt"
        output.headers["Content-type"] = "text/plain; charset=utf-8"
        return output

    def _export_vuln_docx(self, vuln_data: Dict, filename: str) -> Response:
        """Exporta vulnerabilidade para DOCX no formato padronizado"""
        if not HAS_DOCX:
            logging.warning("python-docx não disponível, usando formato TXT como fallback para DOCX")
            return self._export_vuln_txt(vuln_data, filename.replace('.docx', '.txt'))

        doc = Document()
        cve = vuln_data.get('cve', {})
        cve_id = cve.get('id', 'CVE Desconhecido')

        # Cabeçalho BNVD
        heading = doc.add_heading('BNVD - Banco Nacional de Vulnerabilidades Cibernéticas', 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # CVE ID
        title = doc.add_heading(cve_id, 1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Dados básicos
        doc.add_paragraph(f"Status: {cve.get('vulnStatus', 'N/A')}")
        doc.add_paragraph(f"Publicado: {cve.get('published', 'N/A')}")
        doc.add_paragraph(f"Última Modificação: {cve.get('lastModified', 'N/A')}")
        doc.add_paragraph(f"Fonte: {cve.get('sourceIdentifier', 'N/A')}")
        doc.add_paragraph('')

        # Descrição original
        doc.add_heading('DESCRIÇÃO ORIGINAL (INGLÊS)', 2)
        descriptions = cve.get('descriptions', [])
        for desc in descriptions:
            if desc.get('lang') == 'en':
                doc.add_paragraph(desc.get('value', 'N/A'))
                break

        # Descrição traduzida
        doc.add_heading('DESCRIÇÃO TRADUZIDA (PORTUGUÊS)', 2)
        for desc in descriptions:
            if desc.get('lang') == 'en' and desc.get('value_pt'):
                doc.add_paragraph(desc.get('value_pt', 'N/A'))
                break

        # CVSS Metrics detalhado
        doc.add_heading('MÉTRICAS CVSS DETALHADAS', 2)
        metrics = cve.get('metrics', {})
        for version in ['cvssMetricV31', 'cvssMetricV30', 'cvssMetricV2']:
            if version in metrics and metrics[version]:
                metric = metrics[version][0]
                cvss_data = metric.get('cvssData', {})
                doc.add_paragraph(f"Versão: {version}")
                doc.add_paragraph(f"Vector String: {cvss_data.get('vectorString', 'N/A')}")
                doc.add_paragraph(f"Score Base: {cvss_data.get('baseScore', 'N/A')}")
                doc.add_paragraph(f"Severidade: {cvss_data.get('baseSeverity', 'N/A')}")
                doc.add_paragraph(f"Vetor de Ataque: {cvss_data.get('attackVector', 'N/A')}")
                doc.add_paragraph(f"Complexidade: {cvss_data.get('attackComplexity', 'N/A')}")
                doc.add_paragraph(f"Privilégios Necessários: {cvss_data.get('privilegesRequired', 'N/A')}")
                doc.add_paragraph(f"Interação do Usuário: {cvss_data.get('userInteraction', 'N/A')}")
                doc.add_paragraph(f"Escopo: {cvss_data.get('scope', 'N/A')}")
                doc.add_paragraph(f"Impacto Confidencialidade: {cvss_data.get('confidentialityImpact', 'N/A')}")
                doc.add_paragraph(f"Impacto Integridade: {cvss_data.get('integrityImpact', 'N/A')}")
                doc.add_paragraph(f"Impacto Disponibilidade: {cvss_data.get('availabilityImpact', 'N/A')}")
                break

        # CWE
        weaknesses = cve.get('weaknesses', [])
        if weaknesses:
            doc.add_heading('TIPOS DE FRAQUEZA (CWE)', 2)
            for weakness in weaknesses:
                for desc in weakness.get('description', []):
                    doc.add_paragraph(f"• {desc.get('value', 'N/A')}")

        # Referências
        references = cve.get('references', [])
        if references:
            doc.add_heading('REFERÊNCIAS', 2)
            for ref in references:
                doc.add_paragraph(f"URL: {ref.get('url', 'N/A')}")
                doc.add_paragraph(f"Fonte: {ref.get('source', 'N/A')}")
                tags = ref.get('tags', [])
                if tags:
                    doc.add_paragraph(f"Tags: {', '.join(tags)}")
                doc.add_paragraph('')

        doc.add_paragraph('')
        doc.add_paragraph(f"Exportado em: {get_brazil_time().strftime('%d/%m/%Y às %H:%M:%S')} (Horário de Brasília)")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        output = make_response(buffer.getvalue())
        output.headers["Content-Disposition"] = f"attachment; filename={filename}.docx"
        output.headers["Content-type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return output

    def _export_vuln_xml(self, vuln_data: Dict, filename: str) -> Response:
        """Exporta vulnerabilidade para XML no formato padronizado"""
        cve = vuln_data.get('cve', {})
        cve_id = cve.get('id', 'CVE Desconhecido')

        root = ET.Element('bnvd_vulnerability')
        root.set('xmlns', 'http://bnvd.org/vulnerability/1.0')

        # Cabeçalho
        header = ET.SubElement(root, 'header')
        ET.SubElement(header, 'source').text = 'BNVD - Banco Nacional de Vulnerabilidades Cibernéticas'
        ET.SubElement(header, 'type').text = 'Vulnerabilidade'

        # Dados básicos
        basic = ET.SubElement(root, 'basic_info')
        ET.SubElement(basic, 'cve_id').text = cve_id
        ET.SubElement(basic, 'status').text = cve.get('vulnStatus', 'N/A')
        ET.SubElement(basic, 'published').text = cve.get('published', 'N/A')
        ET.SubElement(basic, 'last_modified').text = cve.get('lastModified', 'N/A')
        ET.SubElement(basic, 'source').text = cve.get('sourceIdentifier', 'N/A')

        # Descrições
        descriptions_elem = ET.SubElement(root, 'descriptions')
        desc_original = ET.SubElement(descriptions_elem, 'description_original')
        desc_translated = ET.SubElement(descriptions_elem, 'description_translated')

        for desc in cve.get('descriptions', []):
            if desc.get('lang') == 'en':
                desc_original.text = desc.get('value', 'N/A')
                desc_translated.text = desc.get('value_pt', desc.get('value', 'N/A'))
                break

        # CVSS Metrics detalhado
        metrics_elem = ET.SubElement(root, 'cvss_metrics')
        metrics = cve.get('metrics', {})
        for version in ['cvssMetricV31', 'cvssMetricV30', 'cvssMetricV2']:
            if version in metrics and metrics[version]:
                metric = metrics[version][0]
                cvss_data = metric.get('cvssData', {})
                cvss = ET.SubElement(metrics_elem, 'cvss')
                ET.SubElement(cvss, 'version').text = version
                ET.SubElement(cvss, 'vector_string').text = cvss_data.get('vectorString', 'N/A')
                ET.SubElement(cvss, 'base_score').text = str(cvss_data.get('baseScore', 'N/A'))
                ET.SubElement(cvss, 'severity').text = cvss_data.get('baseSeverity', 'N/A')
                ET.SubElement(cvss, 'attack_vector').text = cvss_data.get('attackVector', 'N/A')
                ET.SubElement(cvss, 'attack_complexity').text = cvss_data.get('attackComplexity', 'N/A')
                ET.SubElement(cvss, 'privileges_required').text = cvss_data.get('privilegesRequired', 'N/A')
                ET.SubElement(cvss, 'user_interaction').text = cvss_data.get('userInteraction', 'N/A')
                ET.SubElement(cvss, 'scope').text = cvss_data.get('scope', 'N/A')
                ET.SubElement(cvss, 'confidentiality_impact').text = cvss_data.get('confidentialityImpact', 'N/A')
                ET.SubElement(cvss, 'integrity_impact').text = cvss_data.get('integrityImpact', 'N/A')
                ET.SubElement(cvss, 'availability_impact').text = cvss_data.get('availabilityImpact', 'N/A')
                break

        # CWE
        weaknesses_elem = ET.SubElement(root, 'weaknesses')
        weaknesses = cve.get('weaknesses', [])
        for weakness in weaknesses:
            for desc in weakness.get('description', []):
                ET.SubElement(weaknesses_elem, 'cwe').text = desc.get('value', 'N/A')

        # Referências
        refs_elem = ET.SubElement(root, 'references')
        for ref in cve.get('references', []):
            ref_elem = ET.SubElement(refs_elem, 'reference')
            ET.SubElement(ref_elem, 'url').text = ref.get('url', 'N/A')
            ET.SubElement(ref_elem, 'source').text = ref.get('source', 'N/A')
            tags = ref.get('tags', [])
            if tags:
                ET.SubElement(ref_elem, 'tags').text = ', '.join(tags)

        # Exportação
        ET.SubElement(root, 'exported_at').text = get_brazil_time().strftime('%d/%m/%Y às %H:%M:%S') + ' (Horário de Brasília)'

        xml_str = minidom.parseString(ET.tostring(root)).toprettyxml(indent="  ")

        output = make_response(xml_str)
        output.headers["Content-Disposition"] = f"attachment; filename={filename}.xml"
        output.headers["Content-type"] = "application/xml; charset=utf-8"
        return output

    def _export_vuln_odf(self, vuln_data: Dict, filename: str) -> Response:
        """Exporta vulnerabilidade para ODF no formato padronizado"""
        cve = vuln_data.get('cve', {})
        cve_id = cve.get('id', 'CVE Desconhecido')

        content = f"""<?xml version="1.0" encoding="UTF-8"?>
<office:document-content xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" 
                         xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0">
  <office:body>
    <office:text>
      <text:h text:style-name="Heading_20_1">BNVD - Banco Nacional de Vulnerabilidades Cibernéticas</text:h>
      <text:h text:style-name="Heading_20_2">{cve_id}</text:h>
      <text:p>Status: {cve.get('vulnStatus', 'N/A')}</text:p>
      <text:p>Publicado: {cve.get('published', 'N/A')}</text:p>
      <text:p>Última Modificação: {cve.get('lastModified', 'N/A')}</text:p>
      <text:p>Fonte: {cve.get('sourceIdentifier', 'N/A')}</text:p>
      <text:p></text:p>
      <text:h text:style-name="Heading_20_3">DESCRIÇÃO ORIGINAL (INGLÊS)</text:h>
"""

        descriptions = cve.get('descriptions', [])
        for desc in descriptions:
            if desc.get('lang') == 'en':
                content += f"      <text:p>{desc.get('value', 'N/A')}</text:p>\n"
                break

        content += "      <text:p></text:p>\n      <text:h text:style-name=\"Heading_20_3\">DESCRIÇÃO TRADUZIDA (PORTUGUÊS)</text:h>\n"

        for desc in descriptions:
            if desc.get('lang') == 'en' and desc.get('value_pt'):
                content += f"      <text:p>{desc.get('value_pt', 'N/A')}</text:p>\n"
                break

        content += "      <text:p></text:p>\n      <text:h text:style-name=\"Heading_20_3\">MÉTRICAS CVSS DETALHADAS</text:h>\n"

        metrics = cve.get('metrics', {})
        for version in ['cvssMetricV31', 'cvssMetricV30', 'cvssMetricV2']:
            if version in metrics and metrics[version]:
                metric = metrics[version][0]
                cvss_data = metric.get('cvssData', {})
                content += f"      <text:p>Versão: {version}</text:p>\n"
                content += f"      <text:p>Vector String: {cvss_data.get('vectorString', 'N/A')}</text:p>\n"
                content += f"      <text:p>Score Base: {cvss_data.get('baseScore', 'N/A')}</text:p>\n"
                content += f"      <text:p>Severidade: {cvss_data.get('baseSeverity', 'N/A')}</text:p>\n"
                content += f"      <text:p>Vetor de Ataque: {cvss_data.get('attackVector', 'N/A')}</text:p>\n"
                content += f"      <text:p>Complexidade: {cvss_data.get('attackComplexity', 'N/A')}</text:p>\n"
                content += f"      <text:p>Privilégios Necessários: {cvss_data.get('privilegesRequired', 'N/A')}</text:p>\n"
                content += f"      <text:p>Interação do Usuário: {cvss_data.get('userInteraction', 'N/A')}</text:p>\n"
                content += f"      <text:p>Escopo: {cvss_data.get('scope', 'N/A')}</text:p>\n"
                content += f"      <text:p>Impacto Confidencialidade: {cvss_data.get('confidentialityImpact', 'N/A')}</text:p>\n"
                content += f"      <text:p>Impacto Integridade: {cvss_data.get('integrityImpact', 'N/A')}</text:p>\n"
                content += f"      <text:p>Impacto Disponibilidade: {cvss_data.get('availabilityImpact', 'N/A')}</text:p>\n"
                break

        weaknesses = cve.get('weaknesses', [])
        if weaknesses:
            content += "      <text:p></text:p>\n      <text:h text:style-name=\"Heading_20_3\">TIPOS DE FRAQUEZA (CWE)</text:h>\n"
            for weakness in weaknesses:
                for desc in weakness.get('description', []):
                    content += f"      <text:p>• {desc.get('value', 'N/A')}</text:p>\n"

        references = cve.get('references', [])
        if references:
            content += "      <text:p></text:p>\n      <text:h text:style-name=\"Heading_20_3\">REFERÊNCIAS</text:h>\n"
            for ref in references:
                content += f"      <text:p>URL: {ref.get('url', 'N/A')}</text:p>\n"
                content += f"      <text:p>Fonte: {ref.get('source', 'N/A')}</text:p>\n"
                tags = ref.get('tags', [])
                if tags:
                    content += f"      <text:p>Tags: {', '.join(tags)}</text:p>\n"
                content += "      <text:p></text:p>\n"

        content += f"      <text:p>Exportado em: {get_brazil_time().strftime('%d/%m/%Y às %H:%M:%S')} (Horário de Brasília)</text:p>\n"
        content += """    </office:text>
  </office:body>
</office:document-content>"""

        output = make_response(content)
        output.headers["Content-Disposition"] = f"attachment; filename={filename}.odt"
        output.headers["Content-type"] = "application/vnd.oasis.opendocument.text"
        return output

    # ==================== Exportação de Notícias ====================

    def _export_news_csv(self, news_data: Dict, filename: str) -> Response:
        """Exporta notícia para CSV com todos os campos no formato padronizado"""
        si = StringIO()
        writer = csv.writer(si)

        # Cabeçalho
        writer.writerow(['BNVD - Notícias de Segurança Cibernética'])
        writer.writerow([''])
        writer.writerow(['Campo', 'Valor'])
        writer.writerow([''])

        # Dados da notícia
        writer.writerow(['Título', news_data.get('title', 'N/A')])
        writer.writerow(['Slug', news_data.get('slug', 'N/A')])
        writer.writerow(['Data de Publicação', news_data.get('pub_date_formatted', news_data.get('pub_date', 'N/A'))])
        writer.writerow(['Categoria', news_data.get('category', 'N/A')])
        writer.writerow(['Link Original', news_data.get('link', 'N/A')])
        writer.writerow(['Autor', 'bnvd.org fonte e autor original cisoadvisor.com.br'])
        writer.writerow([''])

        # Resumo
        writer.writerow(['--- RESUMO/DESCRIÇÃO ---'])
        writer.writerow(['', news_data.get('description', 'N/A')])
        writer.writerow([''])

        # Conteúdo completo
        writer.writerow(['--- CONTEÚDO COMPLETO ---'])
        writer.writerow(['', news_data.get('content', news_data.get('description', 'N/A'))])
        writer.writerow([''])

        writer.writerow(['Exportado em', get_brazil_time().strftime('%d/%m/%Y às %H:%M:%S') + ' (Horário de Brasília)'])

        output = make_response(si.getvalue())
        output.headers["Content-Disposition"] = f"attachment; filename={filename}.csv"
        output.headers["Content-type"] = "text/csv; charset=utf-8"
        return output

    def _export_news_json(self, news_data: Dict, filename: str) -> Response:
        """Exporta notícia para JSON com todos os campos padronizados"""
        export_data = {
            'bnvd': 'Banco Nacional de Vulnerabilidades Cibernéticas',
            'tipo': 'Notícia de Segurança Cibernética',
            'dados': {
                'titulo': news_data.get('title', 'N/A'),
                'slug': news_data.get('slug', 'N/A'),
                'data_publicacao': news_data.get('pub_date_formatted', news_data.get('pub_date', 'N/A')),
                'categoria': news_data.get('category', 'N/A'),
                'link_original': news_data.get('link', 'N/A'),
                'autor': 'bnvd.org fonte e autor original cisoadvisor.com.br',
                'resumo': news_data.get('description', 'N/A'),
                'conteudo_completo': news_data.get('content', news_data.get('description', 'N/A'))
            },
            'exportado_em': get_brazil_time().strftime('%d/%m/%Y às %H:%M:%S') + ' (Horário de Brasília)'
        }

        output = make_response(json.dumps(export_data, indent=2, ensure_ascii=False))
        output.headers["Content-Disposition"] = f"attachment; filename={filename}.json"
        output.headers["Content-type"] = "application/json; charset=utf-8"
        return output

    def _export_news_pdf(self, news_data: Dict, filename: str) -> Response:
        """Exporta notícia para PDF no formato padronizado"""
        if not HAS_REPORTLAB:
            logging.warning("ReportLab não disponível, usando formato TXT como fallback para PDF")
            return self._export_news_txt(news_data, filename.replace('.pdf', '.txt'))

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        story = []
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=20, textColor=colors.HexColor('#0d6efd'), spaceAfter=20, alignment=TA_CENTER)
        heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading2'], fontSize=14, textColor=colors.HexColor('#0d6efd'), spaceAfter=10)

        # Cabeçalho BNVD
        story.append(Paragraph('BNVD - Notícias de Segurança Cibernética', heading_style))
        story.append(Spacer(1, 12))
        story.append(Paragraph(news_data.get('title', 'Notícia'), title_style))
        story.append(Spacer(1, 12))

        # Dados básicos
        story.append(Paragraph(f"<b>Slug:</b> {news_data.get('slug', 'N/A')}", styles['Normal']))
        story.append(Paragraph(f"<b>Data de Publicação:</b> {news_data.get('pub_date_formatted', news_data.get('pub_date', 'N/A'))}", styles['Normal']))
        story.append(Paragraph(f"<b>Categoria:</b> {news_data.get('category', 'N/A')}", styles['Normal']))
        story.append(Paragraph(f"<b>Autor:</b> bnvd.org fonte e autor original cisoadvisor.com.br", styles['Normal']))
        story.append(Paragraph(f"<b>Link Original:</b> {news_data.get('link', 'N/A')}", styles['Normal']))
        tags = news_data.get('tags', [])
        if tags:
            story.append(Paragraph(f"<b>Tags:</b> {', '.join(tags)}", styles['Normal']))
        story.append(Spacer(1, 20))

        # Resumo
        story.append(Paragraph('RESUMO/DESCRIÇÃO', heading_style))
        story.append(Paragraph(news_data.get('description', 'N/A'), styles['Normal']))
        story.append(Spacer(1, 20))

        # Conteúdo completo
        story.append(Paragraph('CONTEÚDO COMPLETO', heading_style))
        story.append(Paragraph(news_data.get('content', news_data.get('description', 'N/A')), styles['Normal']))

        story.append(Spacer(1, 20))
        story.append(Paragraph(f"Exportado em: {get_brazil_time().strftime('%d/%m/%Y às %H:%M:%S')} (Horário de Brasília)", styles['Normal']))

        doc.build(story)
        pdf_data = buffer.getvalue()
        buffer.close()

        output = make_response(pdf_data)
        output.headers["Content-Disposition"] = f"attachment; filename={filename}.pdf"
        output.headers["Content-type"] = "application/pdf"
        return output

    def _export_news_txt(self, news_data: Dict, filename: str) -> Response:
        """Exporta notícia para TXT com todos os campos"""
        content = []
        content.append("=" * 80)
        content.append("BNVD - Notícias de Segurança Cibernética")
        content.append("=" * 80)
        content.append("")
        content.append(news_data.get('title', 'Notícia'))
        content.append("")
        content.append(f"Slug: {news_data.get('slug', 'N/A')}")
        content.append(f"Data de Publicação: {news_data.get('pub_date', 'N/A')}")
        content.append(f"Data de Publicação (UTC): {news_data.get('pub_date_utc', 'N/A')}")
        content.append(f"Categoria: {news_data.get('category', 'N/A')}")
        content.append(f"Autor: bnvd.org fonte e autor original cisoadvisor.com.br")
        content.append(f"Link Original: {news_data.get('link', 'N/A')}")
        content.append(f"Link BNVD: {news_data.get('bnvd_url', 'N/A')}")

        tags = news_data.get('tags', [])
        if tags:
            content.append(f"Tags: {', '.join(tags)}")

        content.append("")
        content.append("-" * 80)
        content.append("RESUMO/DESCRIÇÃO")
        content.append("-" * 80)
        content.append(news_data.get('description', 'N/A'))

        content.append("")
        content.append("-" * 80)
        content.append("CONTEÚDO COMPLETO")
        content.append("-" * 80)
        content.append(news_data.get('content', news_data.get('description', 'N/A')))

        content.append("")
        content.append("=" * 80)
        content.append(f"Exportado em: {get_brazil_time().strftime('%d/%m/%Y às %H:%M:%S')} (Horário de Brasília)")
        content.append("=" * 80)

        txt_data = "\n".join(content)
        output = make_response(txt_data)
        output.headers["Content-Disposition"] = f"attachment; filename={filename}.txt"
        output.headers["Content-type"] = "text/plain; charset=utf-8"
        return output

    def _export_news_docx(self, news_data: Dict, filename: str) -> Response:
        """Exporta notícia para DOCX no formato padronizado"""
        if not HAS_DOCX:
            logging.warning("python-docx não disponível, usando formato TXT como fallback para DOCX")
            return self._export_news_txt(news_data, filename.replace('.docx', '.txt'))

        doc = Document()

        # Cabeçalho BNVD
        heading = doc.add_heading('BNVD - Notícias de Segurança Cibernética', 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Título
        title = doc.add_heading(news_data.get('title', 'Notícia'), 1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Dados básicos
        doc.add_paragraph(f"Slug: {news_data.get('slug', 'N/A')}")
        doc.add_paragraph(f"Data de Publicação: {news_data.get('pub_date_formatted', news_data.get('pub_date', 'N/A'))}")
        doc.add_paragraph(f"Categoria: {news_data.get('category', 'N/A')}")
        doc.add_paragraph(f"Autor: bnvd.org fonte e autor original cisoadvisor.com.br")
        doc.add_paragraph(f"Link Original: {news_data.get('link', 'N/A')}")
        tags = news_data.get('tags', [])
        if tags:
            doc.add_paragraph(f"Tags: {', '.join(tags)}")
        doc.add_paragraph('')

        # Resumo
        doc.add_heading('RESUMO/DESCRIÇÃO', 2)
        doc.add_paragraph(news_data.get('description', 'N/A'))
        doc.add_paragraph('')

        # Conteúdo completo
        doc.add_heading('CONTEÚDO COMPLETO', 2)
        doc.add_paragraph(news_data.get('content', news_data.get('description', 'N/A')))
        doc.add_paragraph('')

        doc.add_paragraph(f"Exportado em: {get_brazil_time().strftime('%d/%m/%Y às %H:%M:%S')} (Horário de Brasília)")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        output = make_response(buffer.getvalue())
        output.headers["Content-Disposition"] = f"attachment; filename={filename}.docx"
        output.headers["Content-type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return output

    def _export_news_xml(self, news_data: Dict, filename: str) -> Response:
        """Exporta notícia para XML com formato padronizado"""
        root = ET.Element('noticia')
        root.set('xmlns', 'http://bnvd.org/news/1.0')

        # Metadados
        meta = ET.SubElement(root, 'metadados')

        title_elem = ET.SubElement(meta, 'titulo')
        title_elem.text = news_data.get('title', 'N/A')

        slug_elem = ET.SubElement(meta, 'slug')
        slug_elem.text = news_data.get('slug', 'N/A')

        date_elem = ET.SubElement(meta, 'data_publicacao')
        date_elem.text = news_data.get('pub_date_formatted', news_data.get('pub_date', 'N/A'))

        category_elem = ET.SubElement(meta, 'categoria')
        category_elem.text = news_data.get('category', 'N/A')

        link_elem = ET.SubElement(meta, 'link_original')
        link_elem.text = news_data.get('link', 'N/A')

        author_elem = ET.SubElement(meta, 'autor')
        author_elem.text = 'bnvd.org fonte e autor original cisoadvisor.com.br'

        # Conteúdo
        desc_elem = ET.SubElement(root, 'resumo')
        desc_elem.text = news_data.get('description', 'N/A')

        content_elem = ET.SubElement(root, 'conteudo_completo')
        content_elem.text = news_data.get('content', news_data.get('description', 'N/A'))

        # Exportação
        export_elem = ET.SubElement(root, 'exportado_em')
        export_elem.text = get_brazil_time().strftime('%d/%m/%Y às %H:%M:%S') + ' (Horário de Brasília)'

        xml_str = minidom.parseString(ET.tostring(root)).toprettyxml(indent="  ")

        output = make_response(xml_str)
        output.headers["Content-Disposition"] = f"attachment; filename={filename}.xml"
        output.headers["Content-type"] = "application/xml; charset=utf-8"
        return output

    def _export_news_odf(self, news_data: Dict, filename: str) -> Response:
        """Exporta notícia para ODF no formato padronizado"""
        content = f"""<?xml version="1.0" encoding="UTF-8"?>
<office:document-content xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" 
                         xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0">
  <office:body>
    <office:text>
      <text:h text:style-name="Heading_20_1">BNVD - Notícias de Segurança Cibernética</text:h>
      <text:h text:style-name="Heading_20_2">{news_data.get('title', 'Notícia')}</text:h>
      <text:p>Slug: {news_data.get('slug', 'N/A')}</text:p>
      <text:p>Data de Publicação: {news_data.get('pub_date_formatted', news_data.get('pub_date', 'N/A'))}</text:p>
      <text:p>Categoria: {news_data.get('category', 'N/A')}</text:p>
      <text:p>Autor: bnvd.org fonte e autor original cisoadvisor.com.br</text:p>
      <text:p>Link Original: {news_data.get('link', 'N/A')}</text:p>"""

        tags = news_data.get('tags', [])
        if tags:
            content += f"\n      <text:p>Tags: {', '.join(tags)}</text:p>"

        content += f"""
      <text:p></text:p>
      <text:h text:style-name="Heading_20_3">RESUMO/DESCRIÇÃO</text:h>
      <text:p>{news_data.get('description', 'N/A')}</text:p>
      <text:p></text:p>
      <text:h text:style-name="Heading_20_3">CONTEÚDO COMPLETO</text:h>
      <text:p>{news_data.get('content', news_data.get('description', 'N/A'))}</text:p>
      <text:p></text:p>
      <text:p>Exportado em: {get_brazil_time().strftime('%d/%m/%Y às %H:%M:%S')} (Horário de Brasília)</text:p>
    </office:text>
  </office:body>
</office:document-content>"""

        output = make_response(content)
        output.headers["Content-Disposition"] = f"attachment; filename={filename}.odt"
        output.headers["Content-type"] = "application/vnd.oasis.opendocument.text"
        return output


# Instância global do gerenciador de exportação
export_manager = ExportManager()